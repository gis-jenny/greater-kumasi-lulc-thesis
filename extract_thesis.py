"""Extract Methodology chapter from thesis.docx for audit (UTF-8 safe)."""
import io
import sys
from docx import Document

# Force UTF-8 output for Windows console
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

path = r"C:\Users\HP\Desktop\project_thesis\thesis.docx"
doc = Document(path)

# Write everything to a file too, for safety
out_path = r"C:\Users\HP\Desktop\project_thesis\thesis_extracted.txt"
with open(out_path, "w", encoding="utf-8") as f:
    f.write("=" * 80 + "\n")
    f.write("FULL THESIS TEXT (paragraph by paragraph)\n")
    f.write("=" * 80 + "\n")
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        style = para.style.name if para.style else "Normal"
        if text.strip():
            f.write(f"[{i:3d}] ({style}) {text}\n")
        else:
            f.write(f"[{i:3d}] (empty)\n")
    f.write(f"\nTOTAL PARAGRAPHS: {len(doc.paragraphs)}\n")
    f.write("\n=== ALL HEADINGS ===\n")
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else "Normal"
        if "Heading" in style or "Title" in style:
            f.write(f"[{i:3d}] ({style}) {para.text}\n")

print(f"Written to {out_path}")
print(f"Total paragraphs: {len(doc.paragraphs)}")