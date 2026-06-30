"""
Targeted fix-up script after the main rewrite.

Fixes:
1. Delete the stray '1.3 Prior Works' heading (no body — body was deleted but
   heading itself remained in old slot).
2. Reorder Ch.3 sections to 3.1 -> 3.7 (currently 3.7 -> 3.1 due to
   reverse-iteration insert-before logic).
3. Reset the 'Research Approach and Materials' replacement paragraph to Normal
   style (currently Heading 2 because set_text didn't change style).
4. Delete stray Study Area heading left over from old Ch.3.
5. Make sure the 'Chapter 3 intro' Study Area paragraph also exists in the new
   number order (it does, 3.1 at end).
"""

import io
import sys
import copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt

PATH = r"C:\Users\HP\Desktop\project_thesis\thesis.docx"
doc = Document(PATH)


def find_idx(predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text):
            return i
    return None


def delete(p):
    parent = p._element.getparent()
    if parent is not None:
        parent.remove(p._element)


# ---------- Fix 1: delete stray '1.3 Prior Works' heading ----------------
prior_idx = find_idx(lambda t: t.strip() == "1.3 Prior Works")
if prior_idx is not None:
    delete(doc.paragraphs[prior_idx])


# ---------- Fix 3: reset 'Overview of research approach' style to Normal --
overview_idx = find_idx(lambda t: t.startswith("An overview of the research approach"))
if overview_idx is not None:
    p = doc.paragraphs[overview_idx]
    try:
        p.style = doc.styles['Normal']
    except KeyError:
        pass


# ---------- Fix 4: delete stray 'Study Area' heading from old Ch.3 --------
stray_study_idx = find_idx(lambda t: t.strip() == "Study Area" or t.strip() == "Study Area ")
# Need to be careful: '3.1 Study Area' should NOT be deleted. Match only the
# bare 'Study Area' heading.
delete_idx = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Study Area":
        delete_idx = i
        break
if delete_idx is not None:
    delete(doc.paragraphs[delete_idx])


# ---------- Fix 2: reorder Ch.3 sections to 3.1 -> 3.7 ---------------------
# Strategy: collect all paragraphs between "3.1 Study Area" (or "CHAPTER THREE"
# header region) and "References" in their current order; sort by section
# number; re-insert in correct order using a moving cursor.

# Find the start of Chapter 3 substantive content: the '3.1 Study Area' heading
# at the bottom of the current order (which is near the end of doc, just before
# References). Find References.
ref_idx = find_idx(lambda t: t.strip() == "References")
start_idx = find_idx(lambda t: t.strip() == "3.1 Study Area")

# Re-scan for Ch.3 block span using '3.1 Study Area' and 'References' markers.
# Note that currently 3.7 is at top, 3.1 at bottom.
ch3_first_idx = None
ch3_last_idx = None
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else ""
    t = p.text.strip()
    if style.startswith("Heading"):
        # Collect Ch.3 headings only.
        if t.startswith("3.1 ") or t.startswith("3.2 ") or t.startswith("3.3 ") \
           or t.startswith("3.4 ") or t.startswith("3.5 ") or t.startswith("3.6 ") \
           or t.startswith("3.7 ") or t.startswith("3.2."):
            if ch3_first_idx is None:
                ch3_first_idx = i
            ch3_last_idx = i

if ch3_first_idx is None or ch3_last_idx is None or ref_idx is None:
    print("Could not locate Ch.3 range; aborting re-order")
else:
    # Capture all paragraphs from ch3_first_idx up to (but not including) ref_idx
    # Use the live list once, then capture references to elements.
    paras_live = list(doc.paragraphs)
    block_elems = []
    for i in range(ch3_first_idx, ref_idx):
        block_elems.append(paras_live[i]._element)
        paras_live[i]._element.getparent().remove(paras_live[i]._element)
    # Sort by leading section number, then by subnumber
    import re


    def sort_key(elem):
        ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        text = "".join(t.text or "" for t in elem.iter(ns + "t"))
        m = re.match(r"^(\d+)(?:\.(\d+))?", text.strip())
        if m:
            a = int(m.group(1))
            b = int(m.group(2)) if m.group(2) else 0
            return (a, b)
        return (99, 99)


    block_elems.sort(key=sort_key)
    # Re-insert before the References element.
    # Find References paragraph again (it should still be there, just at a
    # possibly different index).
    refs_p = None
    for p in doc.paragraphs:
        if p.text.strip() == "References":
            refs_p = p
            break
    if refs_p is None:
        print("References not found after detach")
    else:
        refs_elem = refs_p._element
        for elem in block_elems:
            refs_elem.addprevious(elem)


# ---------- Save ----------------------------------------------------------
out = r"C:\Users\HP\Desktop\project_thesis\thesis.docx"
doc.save(out)
print(f"Wrote {out}")
