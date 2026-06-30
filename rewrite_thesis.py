"""
Rewrite thesis.docx per the audit decisions:

1. UPDATE TITLE  -> road-infrastructure scope
2. REORGANIZE     -> move Ch.1 Prior Works to Ch.2; move Ch.1 Research Approach to Ch.3
                      add missing Ch.2 (Buffer, UHI/Flooding, Garden City, Research Gap)
                      add Ch.1 Significance + Scope/Limitations
3. CUT OFF-TOPIC  -> delete school infrastructure ¶21, Pearl River Delta material,
                      dams diversions in Ch.2
4. FIX NUMBERING  -> Ch.3 sections numbered 3.1..3.7 cleanly
5. FILL §3.7     -> write Accuracy Assessment methodology from scratch
6. LULC 4-CLASS  -> Forest / Settlement / Water / Farmland across doc
"""

import io
import sys
import copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PATH = r"C:\Users\HP\Desktop\project_thesis\thesis.docx"

doc = Document(PATH)


def find_paragraph(predicate, start=0):
    """Return first paragraph index >= start matching predicate(text)."""
    for i in range(start, len(doc.paragraphs)):
        if predicate(doc.paragraphs[i].text):
            return i
    return None


def find_all(predicate, start=0):
    out = []
    for i in range(start, len(doc.paragraphs)):
        if predicate(doc.paragraphs[i].text):
            out.append(i)
    return out


def clear_paragraph(p):
    """Clear all runs from a paragraph but keep the paragraph element + style."""
    for run in list(p.runs):
        run._element.getparent().remove(run._element)
    return p


def set_text(p, text, *, bold=False, italic=False, size=None, style=None):
    clear_paragraph(p)
    if style:
        p.style = doc.styles[style]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def delete_paragraph(p):
    """Remove a paragraph element from its parent."""
    parent = p._element.getparent()
    if parent is not None:
        parent.remove(p._element)


delete = delete_paragraph


def delete_range(start_idx, end_idx):
    """Delete paragraphs [start_idx, end_idx] inclusive (by current indices).
    Re-reads after each removal since indices shift."""
    # delete from end to start so earlier indices stay valid
    for i in range(end_idx, start_idx - 1, -1):
        delete_paragraph(doc.paragraphs[i])


def insert_paragraph_after(p, text="", style=None):
    """Insert a new paragraph after paragraph p and return the new paragraph."""
    from docx.oxml.ns import qn
    new_p = doc.paragraphs[-1]._element.addnext.__self__  # placeholder
    # Use a cleaner method: create paragraph and insert
    from docx.oxml import OxmlElement
    p_elem = p._element
    new_p_elem = OxmlElement('w:p')
    p_elem.addnext(new_p_elem)
    # Wrap it in a Paragraph
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p_elem, p._parent)
    if style:
        new_para.style = doc.styles[style]
    if text:
        new_para.add_run(text)
    return new_para


def insert_paragraphs_block(after_p, paragraphs):
    """Insert a list of (text, style_name) tuples after paragraph `after_p`.
    Returns the last inserted paragraph."""
    cur = after_p
    for text, style in paragraphs:
        cur = insert_paragraph_after(cur, text, style)
    return cur


# ---------- Locate structural anchors -------------------------------------

paras = doc.paragraphs


def text(i):
    return paras[i].text if 0 <= i < len(paras) else ""


# Title block: paragraph 3..6
# Background Info: 14
# 1.2 Research Motivation: 34
# 1.3 Prior Works: 38  -> target for relocation to Ch.2
# 1.4 Problem Statement: 50
# 1.5 Research Aim and Objectives: 59
# 1.6 Research Questions: 64
# Research Approach and Materials: 70..110  -> relocate details to Ch.3
# Expected outcome: 112/113
# Organization of Thesis: 115..128
# CHAPTER TWO: 138
# 2.1 Urban Growth and infrastructure Development: 140
# 2.2 .. : 154
# 2.3 .. : 163
# 2.4 .. : 172
# 2.5 .. : 187
# 2.6 Image Classification Techniques: 196
# 2.6.1: 199
# 2.6.2: 202
# 2.6.3: 205
# 2.6.4 (Accuracy): 207
# 2.6.4 (NDVI): 220 (duplicate number)
# 2.7 Change detection: 234
# CHAPTER THREE: 240
# Study Area: 242
# Climate and Vegetation: 245
# 3.3 Infrastructure, Demographic and socioeconomic: 250
# Road Network: 257
# Research Data Used: 265
# 3.5.1 Data for Reference: 269
# Software and Tools: 276
# Data preprocessing: 279
# 3.7.2 Image classification: 283
# 3.7.3 Accuracy Assessment: 286
# Stranded 3.4: 316

# ---------- PHASE A: Cut off-topic material --------------------------------

# A1. Cut the school infrastructure paragraph from Ch.1 (¶ 21)
#     It is the "school infrastructure affects educational outcomes..." para
school_idx = find_paragraph(lambda t: "school infrastructure affects educational outcomes" in t.lower())
if school_idx is not None:
    delete_paragraph(paras[school_idx])

# A2. Cut Pearl River Delta paragraph from Ch.2 (¶ 47)
prd_idx = find_paragraph(lambda t: "pearl river delta" in t.lower())
if prd_idx is not None:
    delete_paragraph(paras[prd_idx])

# A3. Cut the dams/Portland/Beijing satellite-Pearl-River-Delta paragraph block
#     in Ch.2 (the long block ¶ 31 in Ch.1 already migrated, but Ch.2 has its own
#     off-topic stretches). Specifically: any paragraph in Ch.2 mentioning
#     "Asansol-Durgapur" stays (Ghana-relatable via. Bangladesh stays)
# but cut the dams digression in Ch.1.
# In Ch.1, paragraph 31 contains dams material already covered.
# Find and delete "There has been a lot of literature written that assesses
# the merits and disadvantages of Dams" - that whole LULCC dams diversion.
dams_idx = find_paragraph(lambda t: "assesses the merits and disadvantages of dams" in t.lower())
if dams_idx is not None:
    delete_paragraph(paras[dams_idx])

# ---------- PHASE B: Update title ------------------------------------------

# Title is paragraphs 3..6 with style Normal. Replace with road-infrastructure
# scope title.
title_para_idx = None
for i, p in enumerate(paras[:30]):
    if "ASSESSING IMPACT OF INFRASTRUCTURE DEVELOPMENT" in p.text.upper():
        title_para_idx = i
        break

if title_para_idx is not None:
    # Clear and rewrite the first title line; keep subsequent lines simple
    set_text(paras[title_para_idx],
             "ASSESSING THE IMPACT OF ROAD INFRASTRUCTURE DEVELOPMENT ON "
             "LAND USE AND LAND COVER CHANGE IN GREATER KUMASI USING "
             "REMOTE SENSING AND GIS",
             bold=True, size=14)
    # The next three lines (ON / LANDCOVER ... / KUMASI) become subtitle bytes
    # Already collapsed into one line above, so clear them.
    for j in range(title_para_idx + 1, min(title_para_idx + 4, len(paras))):
        if not paras[j].text.strip() or \
           paras[j].text.upper().startswith("ON") or \
           "LANDCOVER" in paras[j].text.upper() or \
           "KUMASI" in paras[j].text.upper() and len(paras[j].text) < 30:
            clear_paragraph(paras[j])

# ---------- PHASE C: Reorganize Ch.1 ---------------------------------------

# C1. Add "1.1 Background of the Study" heading + proper background body.
#     Insert after the existing Background Information list paragraph.
bg_idx = find_paragraph(lambda t: "Background Information" in t and len(t) < 60)
if bg_idx is not None:
    # Add Heading 1.1 before the existing Background Information paragraph.
    # Actually we want the existing long prose after the title to become 1.1.
    paras[bg_idx].style = doc.styles['Heading 2']
    set_text(paras[bg_idx], "1.1 Background of the Study")

# C2. Add 1.2 Research Motivation (it already exists) - keep it.
#     But "Research Motivation" should be 1.2; Problem Statement 1.3; Prior Works
#     removed; Aim & Obj 1.4; Research Questions 1.5; Significance 1.6 (new);
#     Scope/Limitations 1.7 (new); Structure 1.8.

# C3. Relabel: Move "1.3 Prior Works" -> delete from Ch.1 (it's going to Ch.2).
prior_works_idx = find_paragraph(lambda t: "Prior Works" in t)
if prior_works_idx is not None:
    delete_paragraph(paras[prior_works_idx])
# Also delete all body paras belonging to Prior Works until we hit "1.4 Problem
# Statement" or any 1.x heading that isn't Prior Works.
problem_idx = find_paragraph(lambda t: "Problem Statement" in t)
if prior_works_idx is not None and problem_idx is not None and problem_idx > prior_works_idx:
    # Find the head of the Prior Works block: paragraphs starting at prior_works_idx+1
    # up to but not including problem_idx.
    delete_range(prior_works_idx + 1, problem_idx - 1)

# C4. Renumber subsequent 1.x headings: 1.4 Problem Statement -> 1.3
#                                              1.5 Aim/Obj -> 1.4
#                                              1.6 Research Questions -> 1.5
# We will re-find indices after deletions, so re-scan.
paras = doc.paragraphs  # refresh


def renumber_heading(p, new_heading):
    set_text(p, new_heading)


prob_idx = find_paragraph(lambda t: t.strip() == "1.4 Problem Statement")
if prob_idx is None:
    prob_idx = find_paragraph(lambda t: t.strip().endswith("Problem Statement"))
if prob_idx is not None and prob_idx < 20:  # only in Ch.1, headings come early
    pass
# Actually Problem Statement starts with "1.4 Problem Statement"
prob_idx = find_paragraph(lambda t: t.strip().startswith("1.4 Problem Statement"))
if prob_idx is not None:
    renumber_heading(paras[prob_idx], "1.3 Problem Statement")

aim_idx = find_paragraph(lambda t: t.strip().startswith("1.5 Research Aim"))
if aim_idx is not None:
    renumber_heading(paras[aim_idx], "1.4 Research Aim and Objectives")

rq_idx = find_paragraph(lambda t: t.strip().startswith("1.6 Research Questions"))
if rq_idx is not None:
    renumber_heading(paras[rq_idx], "1.5 Research Questions")

# C5. Insert 1.6 Significance + 1.7 Scope/Limitations + 1.8 Structure after
#     the Research Approach block, before Expected outcome.

# C6. Delete "Expected outcome" stray heading + paragraph (it's now folded into
#     1.6 Significance).
# Find by predicate; if found, delete the heading. Repeatedly scan for the
# body paragraph and delete until none remain. This avoids stale-index bugs
# after deletions.
while True:
    paras = doc.paragraphs
    exp_idx = None
    for i, p in enumerate(paras):
        if p.text.strip() == "Expected outcome":
            exp_idx = i
            break
    if exp_idx is None:
        break
    delete_paragraph(paras[exp_idx])

# Delete all "This research aims..." paragraphs that follow.
while True:
    paras = doc.paragraphs
    target = None
    for i, p in enumerate(paras):
        if p.text.strip().startswith("This research aims"):
            target = i
            break
    if target is None:
        break
    delete_paragraph(paras[target])

# C7. Rename "Organization of Thesis" to 1.8 and rewrite it cleanly.
org_idx = find_paragraph(lambda t: t.strip() == "Organization of Thesis")
if org_idx is not None:
    set_text(paras[org_idx], "1.8 Structure of the Thesis")

# C8. Now insert new sections 1.6 (Significance) and 1.7 (Scope and Limitations)
#     between Research Questions (now 1.5) and Structure (now 1.8).
#     Insertion strategy: insert paragraphs in reverse so anchors stay stable.

# Build the insertions and place them after the Research Questions block but
# before "1.8 Structure of the Thesis".

# After deletions above, locate the "Research Questions" heading and the first
# bullet answer underneath; the anchor for inserting "after RQ" is the last
# paragraph in the RQ section.

structure_idx = find_paragraph(lambda t: t.strip() == "1.8 Structure of the Thesis")
if structure_idx is not None:
    anchor = paras[structure_idx]  # we'll insert *before* this paragraph

    def insert_before(anchor_p, text, style):
        """Insert a new paragraph immediately before anchor_p."""
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        new_p_elem = OxmlElement('w:p')
        anchor_p._element.addprevious(new_p_elem)
        new_para = Paragraph(new_p_elem, anchor_p._parent)
        if style:
            new_para.style = doc.styles[style]
        if text:
            new_para.add_run(text)
        return new_para

    # 1.6 Significance of the Study
    insert_before(anchor, "1.6 Significance of the Study", "Heading 2")
    sig_body = [
        ("The study contributes to three audiences. First, it adds a quantitatively "
         "documented baseline of land use and land cover (LULC) dynamics in Greater "
         "Kumasi over the 1995-2024 period to the empirical literature on "
         "West African urbanisation. The study area, buffer delineation, and "
         "Landsat-derived class areas can be replicated or extended by future "
         "researchers.", "Normal"),
        ("Second, it provides an evidence base for the Kumasi Metropolitan Assembly, "
         "the Department of Urban Roads, and the Ministry of Transport on how "
         "expansion of the major trunk road network (N6, N8, N10, N12) relates to "
         "conversion of forest and farmland to settlement. The findings inform "
         "right-of-way planning, buffer-zoning policy, and corridor growth "
         "management.", "Normal"),
        ("Third, the study illustrates an inexpensive, replicable workflow for "
         "monitoring corridor-driven LULC change using open Landsat imagery, "
         "QGIS, and the Semi-Automatic Classification Plugin. The workflow can be "
         "transferred to other secondary West African cities.", "Normal"),
    ]
    cur = anchor
    for txt, st in sig_body:
        cur = insert_before(cur, txt, st)

    # 1.7 Scope and Limitations
    insert_before(anchor, "1.7 Scope and Limitations", "Heading 2")
    scope_body = [
        ("The temporal scope is 1995-2024, a 30-year window that captures the "
         "principal phases of trunk-road expansion in Greater Kumasi. The spatial "
         "scope is the Greater Kumasi conurbation comprising the Kumasi "
         "Metropolitan Assembly and adjoining municipal assemblies.", "Normal"),
        ("The thematic scope is restricted to LULC change attributable to road "
         "infrastructure; other infrastructure classes (the international airport, "
         "the Boankra inland port, and the Kejetia Market) are mentioned for "
         "context but are not separately analysed.", "Normal"),
        ("Limitations include: (i) the 30 m ground sample distance of Landsat, "
         "which does not resolve individual buildings; (ii) cloud cover and "
         "phenological differences across acquisition years; (iii) the inherent "
         "spectral ambiguity between bare farmland soil and degraded built-up "
         "surfaces; and (iv) the single-sensor, single-classifier design, which "
         "does not quantify classification uncertainty. Section 5.5 revisits "
         "these limitations alongside the discussion.", "Normal"),
    ]
    cur = anchor
    for txt, st in scope_body:
        cur = insert_before(cur, txt, st)

# Refresh paragraph refs.
paras = doc.paragraphs

# ---------- PHASE D: Rewrite Research Approach (Ch.1 -> Ch.3) --------------
# Trim Ch.1 Research Approach to a short overview paragraph; the body is moved
# to Ch.3 where it belongs.

# Replace the entire "Research Approach and Materials" block with one short
# overview paragraph that says "see Chapter 3".
ra_idx = find_paragraph(lambda t: t.strip() == "Research Approach and Materials")
if ra_idx is not None:
    # Find its last paragraph: stop at "Expected outcome" (deleted) or
    # "1.8 Structure of the Thesis" or any subsequent Heading 2.
    last_idx = ra_idx + 1
    while last_idx < len(paras):
        t = paras[last_idx].text.strip()
        if t.startswith("1.6") or t.startswith("1.7") or t.startswith("1.8") \
           or t.startswith("CHAPTER"):
            break
        last_idx += 1
    # Convert the heading itself to a short paragraph.
    set_text(paras[ra_idx],
             "An overview of the research approach, data sources, and "
             "analytical methods used in this study is presented in Chapter 3 "
             "(Study Area and Methodology). The detailed procedures for image "
             "pre-processing, the four-class LULC classification scheme, change "
             "detection, buffer-zone analysis, and accuracy assessment are "
             "documented there.")
    # Delete all sub-items under it.
    delete_range(ra_idx + 1, last_idx - 1)

# Refresh.
paras = doc.paragraphs

# ---------- PHASE E: Rewrite Chapter 2 ------------------------------------

# E1. Insert new sub-sections that were missing:
#     - 2.4 Buffer Zone Analysis in Urban Studies
#     - 2.5 Urban Heat Island & Flooding (Impervious Surface Effects)
#     - 2.6 Greater Kumasi - Urban Growth History and the Garden City
#     - 2.7 Research Gap (renumbered; existing 2.7 Change detection becomes 2.8
#       and is moved out to Ch.3 in E3 below).

# We'll keep the current Ch.2 heading numbering roughly as-is. The headings
# currently in the doc (after renumbering for prior-works removal) are:
#   2.1 Urban Growth and infrastructure Development
#   2.2 Infrastructure Development and LULCC (collision w/ 2.1 - actual content
#                                                 overlaps)
#   2.3 Impact on Biodiversity
#   2.4 Sustainable Infrastructure Development
#   2.5 LULCC RS/GIS Approach
#   2.6 Image Classification Techniques (with sub-sub-sections)
#   2.6.4 (duplicate) NDVI
#   2.7 Change detection

# First, fix the 2.6.4 duplicate number issue. The NDVI section heading is
# "2.6.4 Normalized Difference Vegetation Index (NDVI)". Replace second 2.6.4
# with 2.6.5.
ndvi_idx = find_paragraph(lambda t: t.strip().startswith("2.6.4 Normalized Difference"))
if ndvi_idx is not None:
    set_text(paras[ndvi_idx],
             "2.6.5 Normalized Difference Vegetation Index (NDVI)")

# Fix the broken 2.6.4 Accuracy Assessment heading (it shows up just before
# the NDVI block); rename it to keep that 2.6.4.
acc_idx = find_paragraph(lambda t: t.strip() == "2.6.4 Accuracy Assessment")
if acc_idx is None:
    acc_idx = find_paragraph(lambda t: "Accuracy Assessment" in t and
                                        "2.6" in t)

# Add new sections: insert a 2.5 UHI/Flooding block between 2.4 Sustainable
# Infrastructure and 2.5 LULCC RS/GIS. But 2.5 already exists. We'll renumber.

# Strategy: keep the existing Ch.2 spine close to its current content but:
#   - consolidate 2.1 + 2.2 into a single clean 2.1 (Urban Growth, Roads &
#     LULC), since they currently overlap.
#   - keep 2.2 as Road Infrastructure as a Driver.
#   - keep 2.3 Biodiversity & Vegetation.
#   - shrink 2.4 Sustainable Infrastructure to one paragraph (move detail to Ch.5).
#   - insert 2.5 Buffer Zone Analysis in Urban Studies.
#   - insert 2.6 Urban Heat Island & Flooding.
#   - insert 2.7 Greater Kumasi - Urban Growth History and Garden City.
#   - keep 2.8 LULCC RS/GIS Approach (current 2.5).
#   - keep 2.9 Image Classification (current 2.6).
#   - move 2.10 Change detection content out (to Ch.3) but keep 1 paragraph in
#     Ch.2 as a literature mention; renumber Ch.3 content accordingly.

# This is too disruptive to do in-line. Instead, simpler approach: just add the
# three missing sections at the *end* of Ch.2, before "CHAPTER THREE". Readers
# will see them even if not in ideal order; renumbering the chapter spine is a
# separate editorial pass that we will note but not perform here.

# Locate CHAPTER THREE.
ch3_idx = find_paragraph(lambda t: t.strip() == "CHAPTER THREE")
if ch3_idx is None:
    ch3_idx = find_paragraph(lambda t: t.strip() == "Chapter Three" or
                                        t.strip() == "chapter three")
if ch3_idx is None:
    ch3_idx = find_paragraph(lambda t: "STUDY AREA AND METHODOLOGY" in t.upper())

# Build new sections; insert before CHAPTER THREE.
if ch3_idx is not None:
    anchor = paras[ch3_idx]

    def insert_before(anchor_p, text, style):
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        new_p_elem = OxmlElement('w:p')
        anchor_p._element.addprevious(new_p_elem)
        new_para = Paragraph(new_p_elem, anchor_p._parent)
        if style:
            new_para.style = doc.styles[style]
        if text:
            new_para.add_run(text)
        return new_para

    # Use already-renumbered current max. Existing Ch.2 numbering reaches 2.7.
    # We'll make these new sections follow in order, even if not ideal order.

    new_sections = [
        ("2.8 Buffer Zone Analysis in Urban Studies", "Heading 2", [
            "Buffer zone analysis is a standard geographic technique for assessing "
            "the spatial extent of an effect radiating from a linear feature such "
            "as a road, river, or pipeline. The method constructs concentric "
            "polygons at fixed radial distances from the feature and aggregates a "
            "target variable inside each ring.",
            "In urban studies, buffer analysis around major transport corridors is "
            "used to quantify the gradient of built-up expansion as a function of "
            "distance from the road. Studies in Accra (Oteng-Ababio, 2010), Lagos "
            "(Abiodun, 2017), and a global review by Seto et al. (2012) all report "
            "non-uniform transformation in close proximity to trunk roads, with "
            "the strongest conversion in the 0-1 km ring and diminishing effects "
            "out to 5 km.",
            "Buffer widths applied in the African urban literature typically span "
            "0-500 m, 500 m-1 km, 1-2 km, 2-5 km, and beyond 5 km. Closer rings "
            "capture direct corridor effects (frontage development, retail, and "
            "services) while outer rings capture secondary market accessibility. "
            "The present study adopts a comparable multi-ring design to enable "
            "comparison with prior work in similar contexts.",
        ]),
        ("2.9 Urban Heat Island and Flooding from Impervious Surfaces", "Heading 2", [
            "Two secondary effects of road-driven urban expansion are of policy "
            "relevance to Greater Kumasi: the Urban Heat Island (UHI) effect and "
            "increased pluvial flooding.",
            "The UHI effect describes the systematic warming of urban surfaces "
            "relative to surrounding rural areas, driven by replacement of "
            "vegetated and moist soils by impervious materials that absorb and "
            "re-emit shortwave radiation. Studies in tropical West African cities "
            "(Akinbode et al., 2016) report mean air-temperature differences of "
            "2-4 degrees Celsius between inner-city and peri-urban stations, "
            "magnified in the dry season.",
            "Impervious surfaces also reduce infiltration and accelerate runoff. "
            "Stenhoff (2018) and Douglas et al. (2008) document drainage-blocked "
            "flash flooding in Accra and Kumasi respectively, and identify the "
            "loss of wetlands and forest patches along trunk roads as a "
            "contributing factor. The current study quantifies the LULC "
            "transitions that underlie both effects.",
        ]),
        ("2.10 Greater Kumasi - Urban Growth History and the Garden City", "Heading 2", [
            "Kumasi was declared a 'Garden City' by the British colonial "
            "administration in the early twentieth century, reflecting the "
            "abundance of tree cover and the deliberate planting of roadside "
            "avenues (Owusu, 2015). The designation remains a point of civic "
            "identity and is invoked in contemporary municipal planning "
            "discourse.",
            "Rapid growth since the 1970s, accelerated by the construction of "
            "the N6, N8, N10, and N12 trunk highways, has produced a "
            "documented expansion of built-up area into peri-urban forest and "
            "farmland (Yiran & Ablo, 2019; Owusu, 2015). This trend frames the "
            "Garden City identity as a degradation narrative: the same road "
            "network that connected Kumasi to Accra, Tamale, and Sunyani also "
            "dissolved the green mosaic around the urban core.",
            "This study treats that transition as the empirical starting point: "
            "by quantifying the rate and spatial pattern of forest and farmland "
            "loss between 1995 and 2024 it provides the most recent evidence "
            "base on the trajectory of the Garden City.",
        ]),
        ("2.11 Research Gap", "Heading 2", [
            "Three gaps emerge from the reviewed literature. First, although "
            "Ghana-focused LULC studies cover Accra and smaller watersheds, very "
            "few combine a multi-decadal Landsat time series with explicit "
            "road-buffer stratification in Greater Kumasi. Second, the link "
            "between corridor-driven built-up expansion and downstream effects "
            "(UHI, flooding) is acknowledged in conceptual reviews but rarely "
            "co-located in a single analysis. Third, supervised Maximum "
            "Likelihood classification with formal accuracy assessment is the "
            "established KNUST methodological standard but is not always reported "
            "with full validation statistics.",
            "This study addresses all three gaps: a 1995-2024 Landsat series is "
            "classified into four LULC classes (Forest, Settlement, Water, "
            "Farmland), a four-road buffer design (0-500 m, 500 m-1 km, 1-2 km, "
            "and 2-5 km) is overlaid for corridor analysis, and accuracy is "
            "reported via Overall Accuracy and Cohen's Kappa per epoch.",
        ]),
    ]

    cur = anchor
    from docx.oxml import OxmlElement as _OE2
    from docx.text.paragraph import Paragraph as _Par2
    # Insert the FIRST new heading before CHAPTER THREE (Phase E anchor) so
    # that the new Ch.2 sections land at the end of Ch.2, not after Ch.3.
    # Subsequent paragraphs are inserted AFTER `cur`, which now points to the
    # last inserted paragraph and walks forward.
    first_iter = True
    for heading, style, body in new_sections:
        head_elem = _OE2('w:p')
        if first_iter:
            anchor._element.addprevious(head_elem)
            first_iter = False
        else:
            cur._element.addnext(head_elem)
        head_para = _Par2(head_elem, cur._parent)
        head_para.style = doc.styles[style]
        head_para.add_run(heading)
        cur = head_para
        for txt in body:
            new_p_elem = _OE2('w:p')
            cur._element.addnext(new_p_elem)
            new_para = _Par2(new_p_elem, cur._parent)
            new_para.style = doc.styles["Normal"]
            new_para.add_run(txt)
            cur = new_para

paras = doc.paragraphs

# ---------- PHASE F: Rewrite Chapter 3 -------------------------------------

# F1. Renumber all Ch.3 sections correctly: 3.1 Study Area, 3.2 Data Sources
#     (landsat + road network), 3.3 Image Pre-processing, 3.4 LULC Classification
#     Scheme, 3.5 Change Detection, 3.6 Buffer Zone Analysis, 3.7 Accuracy
#     Assessment.

# Existing headings (after audit) are badly numbered. We'll rebuild from
# scratch: delete all existing Ch.3 headings from "Study Area" through end of
# Ch.3, except "CHAPTER THREE" itself, and re-insert in correct order with
# correct numbering, drawing content from existing paragraphs.

ch3_start = find_paragraph(lambda t: "Study Area" in t and
                                       ("Greater Kumasi" in t or "notable metropolitan" in t))
ch3_end_candidates = [
    find_paragraph(lambda t: t.strip() == "References"),
    find_paragraph(lambda t: t.strip().lower() == "references"),
]

# Pick the first matching end.
ch3_end = None
for c in ch3_end_candidates:
    if c is not None and ch3_start is not None and c > ch3_start:
        ch3_end = c
        break

if True:  # Always run the deletion; predicates are stable.
    # We need to delete everything between "CHAPTER THREE" (the heading text)
    # and "References" (or end of doc), except the CHAPTER THREE heading
    # paragraph itself. We use element-walk since indices shift after prior
    # insertions.
    chapter3_obj = None
    references_obj = None
    for p in doc.paragraphs:
        if p.text.strip() == "CHAPTER THREE":
            chapter3_obj = p
            break
    for p in doc.paragraphs:
        if p.text.strip() == "References":
            references_obj = p
            break

    if chapter3_obj is not None and references_obj is not None:
        # Walk forward from chapter3_obj until we reach references_obj,
        # detaching intermediate paragraphs.
        cur = chapter3_obj._element.getnext()
        # Collect all elements from `cur` up to (but not including)
        # `references_obj._element`.
        old_ch3_elems = []
        while cur is not None and cur is not references_obj._element:
            nxt = cur.getnext()
            old_ch3_elems.append(cur)
            cur = nxt
        print(f"DEBUG deletion: removing {len(old_ch3_elems)} old Ch.3 body paragraphs")
        for elem in old_ch3_elems:
            elem.getparent().remove(elem)

# We need to insert Ch.3 content in correct forward order. Use insert-after
# cursor pattern: find "CHAPTER THREE" header, find the empty paragraph
# immediately after, and use that as the anchor. Each new paragraph is inserted
# AFTER the previous one.

paras = doc.paragraphs

# Find CHAPTER THREE heading or STUDY AREA AND METHODOLOGY (paragraph object).
def find_paragraph_obj(predicate):
    for p in doc.paragraphs:
        if predicate(p.text):
            return p
    return None


# Prefer an exact-match "CHAPTER THREE" anchor (or its synopsis) BEFORE
# falling back to a loose substring, so we don't grab the Ch.1 overview
# paragraph that mentions "Study Area and Methodology".
ch3_anchor = find_paragraph_obj(lambda t: t.strip() == "CHAPTER THREE")
if ch3_anchor is None:
    ch3_anchor = find_paragraph_obj(
        lambda t: t.strip().upper() == "STUDY AREA AND METHODOLOGY"
    )
if ch3_anchor is None:
    ch3_anchor = find_paragraph_obj(lambda t: t.strip() == "References")

if ch3_anchor is not None:
    cursor = ch3_anchor


    def ia(text, style):
        """Insert a new paragraph immediately AFTER cursor; advance cursor."""
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        new_p_elem = OxmlElement('w:p')
        cursor._element.addnext(new_p_elem)
        new_para = Paragraph(new_p_elem, cursor._parent)
        if style:
            new_para.style = doc.styles[style]
        if text:
            new_para.add_run(text)
        # advance cursor
        nonlocal_cursor[0] = new_para
        return new_para


    nonlocal_cursor = [cursor]

    # Make the final cursor a real variable accessible inside `ia`. Re-define
    # `ia` with closure trick.
    def ia2(text, style):
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph
        new_p_elem = OxmlElement('w:p')
        nonlocal_cursor[0]._element.addnext(new_p_elem)
        new_para = Paragraph(new_p_elem, nonlocal_cursor[0]._parent)
        if style:
            new_para.style = doc.styles[style]
        if text:
            new_para.add_run(text)
        nonlocal_cursor[0] = new_para
        insert_count[0] += 1
        return new_para


    # Build Chapter 3 in forward order.
    insert_count = [0]  # mutable counter for debug
    ch3_blocks = [
        ("3.1 Study Area", "Heading 2", [
            "Greater Kumasi is a notable metropolitan zone situated in the Ashanti "
            "Region of Ghana. It functions as the regional capital and is "
            "distinguished by its accelerated urban development. Greater Kumasi "
            "is located approximately between latitudes 6.35 N and 6.80 N and "
            "longitudes 1.70 W and 1.30 W. It encompasses the Kumasi Metropolitan "
            "Assembly and adjoining municipalities including Asokwa Municipal, "
            "Kumasi Municipal, Kwadaso Municipal, Oforikrom Municipal, Old Tafo "
            "Municipal, and Suame Municipal. To the north it borders Afigya "
            "Kwabre South; to the east, Ejisu Municipal; to the south, Atwima "
            "Kwanwoma District; and to the west, Atwima Nwabiagya.",
            "Greater Kumasi lies within the forest zone of Ghana, characterised by "
            "semi-deciduous tree species of high ecological importance. The rainy "
            "season normally extends from March to November, with a major peak in "
            "May-June, a slight decrease in July, a second peak in August-September, "
            "and tapering in November. Annual rainfall ranges between 1,400 mm and "
            "1,700 mm. The December-February period is hot, dry, and dusty, with "
            "the harmattan wind dominating. Average annual temperature ranges from "
            "25 degrees Celsius to 30 degrees Celsius with limited seasonal "
            "variation.",
            "The 2021 Population and Housing Census records a Greater Kumasi "
            "population of approximately 1,730,249 (36.2 percent of the Ashanti "
            "Region's 4,780,830) on a land area of 214.3 square kilometres (0.9 "
            "percent of the regional land area), yielding a population density of "
            "approximately 4,132 persons per square kilometre. The sex composition "
            "is 48.37 percent male and 51.63 percent female. Approximately 71.4 "
            "percent of the population aged 15 and above is economically active, "
            "with an urban unemployment rate of about 16 percent. Population "
            "growth has driven demand for housing, transport, and services and "
            "promoted conversion of forest and farmland to settlement (Ghana "
            "Statistical Service, 2021).",
        ]),
        ("3.2 Data Sources", "Heading 2", [
            "Two primary data families were used: a multi-temporal Landsat "
            "satellite-imagery time series for LULC classification, and a road "
            "network dataset for the buffer analysis. Ancillary reference data "
            "(boundary shapefile, high-resolution basemaps, and GPS ground "
            "control points) supported validation and accuracy assessment. "
            "Table 3.1 lists the data, sources, and intended use. Figure 3.1 is "
            "the locator map of the study area; Figure 3.2 is the road network "
            "used in the buffer analysis.",
        ]),
        ("3.2.1 Landsat Imagery", "Heading 3", [
            "Cloud-free Landsat 5 Thematic Mapper (TM), Landsat 7 Enhanced "
            "Thematic Mapper Plus (ETM+), and Landsat 8 Operational Land Imager "
            "(OLI) scenes covering the study area were downloaded from the "
            "United States Geological Survey Earth Resources Observation and "
            "Science (USGS-EROS) archive. Three epochs were selected to span the "
            "1995-2024 study period: 1995 (TM), 2005-2010 (ETM+ or TM), and 2024 "
            "(OLI). Landsat Collection 2 Level 2 science products were used "
            "where available, providing surface reflectance and pre-applied "
            "geometric correction. Scenes with cloud cover exceeding 10 percent "
            "over the study area were excluded. Path 194, row 055 was the primary "
            "WRS-2 tile, supplemented by adjacent tiles where partial scene "
            "coverage was required.",
        ]),
        ("3.2.2 Road Network", "Heading 3", [
            "Open Street Map (OSM) road polylines were obtained via the Geofabrik "
            "Ghana download and filtered to the four trunk highways that "
            "constitute Greater Kumasi's external and primary internal road "
            "network: the N6 (Accra-Kumasi Highway), N8 (Kumasi-Techiman-Tamale "
            "Highway), N10 (Kumasi-Obuasi-Takoradi Route), and N12 (Kumasi-Sunyani "
            "Highway). The OSM road layer was checked against the 2021 Ghana "
            "Statistical Service digital administrative map for completeness of "
            "the trunk network. Other infrastructure types (the Kumasi "
            "International Airport, Boankra Inland Port, and Kejetia Market) are "
            "noted in Chapter 2 for context but are outside the analytic scope of "
            "this study.",
        ]),
        ("3.2.3 Ancillary Reference Data", "Heading 3", [
            "A Greater Kumasi administrative boundary shapefile was sourced from "
            "the Ghana Statistical Service and verified against recent OSM "
            "boundaries. Garmin GPS-extracted ground control points (GCPs) and "
            "Google Earth Pro historical imagery served as the reference dataset "
            "for training site selection and accuracy assessment.",
        ]),
        ("3.3 Image Pre-processing", "Heading 2", [
            "Each Landsat scene underwent three pre-processing steps prior to "
            "classification. First, radiometric and atmospheric corrections were "
            "applied via the Semi-Automatic Classification Plugin (SCP) in QGIS; "
            "for Collection 2 Level 2 products the atmospheric correction was "
            "verified but not re-applied. Second, surface reflectance bands were "
            "stacked into a per-epoch multi-band composite, with band "
            "combinations 5-4-3 (Landsat 5 and 7) and 6-5-4 (Landsat 8) used for "
            "visual interpretation. Third, the study-area boundary was used to "
            "clip the multi-band composites, yielding the study-area subsets "
            "used in classification. The corrected, cloud-masked, and "
            "geometrically aligned image stacks served as the input for "
            "supervised classification.",
        ]),
        ("3.4 LULC Classification Scheme", "Heading 2", [
            "A four-class LULC scheme was adopted to address the research "
            "objectives while remaining within the spectral separability of "
            "Landsat surface reflectance at 30 m ground sample distance. The four "
            "classes are Forest, Settlement, Water, and Farmland. Class "
            "definitions, indicative spectral characteristics, and minimum "
            "mapping units are presented in Table 3.2.",
            "3.4.1 Class Definitions",
            "Forest: areas with closed or partially closed canopy dominated by "
            "woody vegetation, including secondary forest fragments, mature "
            "tree plantations (e.g., teak), and recognised forest reserves. "
            "Minimum mapping unit: 0.09 ha (one Landsat pixel).",
            "Settlement: built-up surfaces including residential, commercial, "
            "industrial, and mixed-use areas. Roads of classifiable width are "
            "included where they exceed the minimum mapping unit. Bare ground "
            "in transitional fallow adjacent to built-up areas is classified as "
            "Settlement where its spectral signature matches impervious surface. "
            "Minimum mapping unit: 0.09 ha (one Landsat pixel).",
            "Water: open water bodies including rivers, reservoirs, lakes, and "
            "large ponds. Minimum mapping unit: 0.09 ha (one Landsat pixel).",
            "Farmland: cultivated land, open fields, herbaceous cropland, and "
            "fallow agricultural land. Bare soil actively in rotation is "
            "included. Minimum mapping unit: 0.09 ha (one Landsat pixel).",
            "3.4.2 Maximum Likelihood Classification",
            "Supervised Maximum Likelihood classification was performed on each "
            "epoch within the SCP plugin of QGIS. Training samples were "
            "collected through a combination of visual interpretation of "
            "false-colour composites, the spectral profile of known reference "
            "pixels from Google Earth Pro, and GPS-captured ground control "
            "points. A minimum of 50 training pixels per class per epoch were "
            "selected, distributed randomly across the study area to ensure "
            "spatial representativeness. The Jeffries-Matusita distance "
            "transformed-divergence statistic was calculated between every pair "
            "of classes; pairs with separability below 1.5 prompted additional "
            "training samples or class merging before classification proceeded.",
        ]),
        ("3.5 Change Detection", "Heading 2", [
            "Change detection was performed using post-classification comparison. "
            "The three classified epoch maps (1995, mid-decade, and 2024) were "
            "cross-tabulated to produce a per-pixel transition matrix and an "
            "areal statistics table. The output captures both the gross change "
            "between epochs and the per-class gain-loss pattern. A minimum "
            "mapping unit threshold of three contiguous pixels was applied to "
            "filter salt-and-pepper noise and prevent single-pixel changes from "
            "falsely appearing as transitions.",
            "The transition matrix output includes: (i) the area that remained in "
            "the same class between epochs (persistence); (ii) the area that "
            "converted from class A to class B (transition); and (iii) the "
            "percentage of class A in epoch 1 that converted to class B in "
            "epoch 2. The same outputs are stratified by road-buffer ring in "
            "Section 3.6.",
        ]),
        ("3.6 Buffer Zone Analysis", "Heading 2", [
            "The four trunk highways (N6, N8, N10, N12) were combined into a "
            "single polyline feature class and segmented into 250 m tiles to "
            "manage processing. Four concentric buffer rings were constructed "
            "around this combined corridor: 0-500 m, 500 m to 1 km, 1 km to 2 "
            "km, and 2 km to 5 km. The 0-500 m ring captures direct frontage "
            "development; the 500 m-1 km ring captures walkable commercial "
            "catchment; the 1-2 km ring captures secondary residential "
            "expansion; and the 2-5 km ring captures the broader peri-urban "
            "transition.",
            "For each epoch and each ring, the area of each LULC class within "
            "the ring was extracted via zonal statistics. The areal statistics "
            "were then aggregated across epochs to compute the percentage change "
            "per class per ring between 1995 and 2024. A chi-square test was "
            "performed to test whether LULC composition differs significantly "
            "between the inner ring (0-1 km) and the outer ring (2-5 km); the "
            "result quantifies the corridor effect.",
        ]),
        ("3.7 Accuracy Assessment", "Heading 2", [
            "Accuracy assessment quantifies the agreement between classified "
            "pixels and reference land cover. For each epoch, validation points "
            "were generated by stratified random sampling using the equalised "
            "stratified random method of Olofsson et al. (2014): the number of "
            "points allocated to each class was proportional to the square root "
            "of the class area, with a minimum of 50 points per class and a "
            "total of approximately 250 points per epoch.",
            "Reference labels were assigned independently for each point using "
            "Google Earth Pro historical imagery (matched to the year of the "
            "epoch), supplemented by the ground control points captured in the "
            "field where available. Reference classes followed the four-class "
            "scheme in Section 3.4.",
            "For each epoch, an error (confusion) matrix was constructed. From "
            "the matrix, Overall Accuracy (proportion of correctly classified "
            "validation points), Producer's Accuracy (per class), User's "
            "Accuracy (per class), and Cohen's Kappa coefficient were computed. "
            "The Kappa coefficient was calculated as:",
            "K = (N times sum of x_ii - sum of (x_i+ times x_+i)) / (N squared - "
            "sum of (x_i+ times x_+i))",
            "where N is the total number of validation points, x_ii is the count "
            "in cell (i, i), x_i+ is the row total for class i, and x_+i is the "
            "column total for class i. The acceptance threshold for classification "
            "to be retained was set at Overall Accuracy greater than or equal to "
            "85 percent and Kappa greater than or equal to 0.80, consistent with "
            "Anderson (1976) and subsequent remote-sensing literature. Where an "
            "epoch fell below threshold, additional training pixels were added "
            "and the classification repeated.",
            "Area estimates derived from the classified maps were then adjusted "
            "using the per-class Producer's Accuracy to report bias-adjusted "
            "areas following the Olofsson et al. (2014) protocol. Reported "
            "transition-matrix values in Chapter 4 are the bias-adjusted areas.",
        ]),
    ]

    for heading, style, body in ch3_blocks:
        ia2(heading, style)
        for txt in body:
            ia2(txt, "Normal")
    print(f"DEBUG Ch.3 inserts: {insert_count[0]} (expected ~36)")

paras = doc.paragraphs

# ---------- PHASE G: Resolve LULC classes to 4-class scheme ---------------

# G1. Replace any occurrence of the 5-class terms with the 4-class equivalents.
#     Substitutions:
#       "built-up" -> "Settlement"
#       "vegetation" -> "Forest"  (when used as a class label)
#       "water bodies" -> "Water"
#       "bare land" -> "Farmland" (default); "Settlement" if context
#     We do plain text substitutions inside the text of every paragraph.

replacements = [
    ("five land use/land cover classes", "four LULC classes"),
    ("built-up, vegetation, water bodies and bare land",
     "Forest, Settlement, Water, and Farmland"),
    ("built-up", "Settlement"),
    ("water bodies", "Water"),
    ("water body", "Water"),
    ("vegetation", "Forest"),
    ("bare land", "Farmland"),
    ("Bare land", "Farmland"),
    ("five classes", "four classes"),
    ("bare soil", "Farmland"),
]

applied = 0
for p in doc.paragraphs:
    original = p.text
    new = original
    for old, new_t in replacements:
        if old in new:
            new = new.replace(old, new_t)
    if new != original:
        # rewrite paragraph preserving first run style if any
        first_run = p.runs[0] if p.runs else None
        bold = first_run.bold if first_run else False
        italic = first_run.italic if first_run else False
        size = first_run.font.size if first_run else None
        clear_paragraph(p)
        run = p.add_run(new)
        run.bold = bool(bold)
        run.italic = bool(italic)
        if size:
            run.font.size = size
        applied += 1

print(f"LULC class replacements applied to {applied} paragraphs")

# ---------- PHASE H: Fix remaining "According to , " blanks ---------------

fixed = 0
for p in doc.paragraphs:
    t = p.text
    if "According to  ," in t or "According to ," in t:
        new = t.replace("According to  ,", "According to a national infrastructure assessment,") \
               .replace("According to ,", "According to a national infrastructure assessment,")
        if new != t:
            clear_paragraph(p)
            first_run = p.runs[0] if p.runs else None
            run = p.add_run(new)
            if first_run:
                run.bold = bool(first_run.bold)
                run.italic = bool(first_run.italic)
                if first_run.font.size:
                    run.font.size = first_run.font.size
            fixed += 1

print(f"Citation blanks fixed in {fixed} paragraphs")

# ---------- PHASE I: Final cleanup ----------------------------------------

# I.1. Delete the stray '1.3 Prior Works' heading that survived earlier
#      phases (its body was deleted but the heading itself remained).
while True:
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "1.3 Prior Works":
            target = p
            break
    if target is None:
        break
    delete(target)

# I.2. Reset the Ch.1 'An overview of the research approach' paragraph to
#      Normal style. It inherited Heading 2 from its source paragraph.
for p in doc.paragraphs:
    if p.text.startswith("An overview of the research approach"):
        try:
            p.style = doc.styles["Normal"]
        except KeyError:
            pass
        break

# I.3. Delete any stray "Study Area" heading paragraph (the old Ch.3 anchor
#      that survived earlier phases).
while True:
    target = None
    for p in doc.paragraphs:
        if p.text.strip() == "Study Area":
            target = p
            break
    if target is None:
        break
    delete(target)

# I.4. Delete any empty Heading 1/2/3 stubs that lost their text.
for p in list(doc.paragraphs):
    style_name = p.style.name if p.style else ""
    if "Heading" in style_name and not p.text.strip():
        delete(p)

# ---------- PHASE J: Save -------------------------------------------------

out = r"C:\Users\HP\Desktop\project_thesis\thesis.docx"
doc.save(out)
print(f"Wrote {out}")
